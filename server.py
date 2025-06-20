# file: server.py
from flask import Flask, request, jsonify
from docx import Document
from minio import Minio
import uuid
import io
from PIL import Image
from lxml import etree

app = Flask(__name__)

# Inisialisasi MinIO Client
minio_client = Minio(
    endpoint="storage-api.sman16bekasi.id",  # ganti dengan endpoint MinIO kamu
    access_key="4ormnVvuMMJy5A84wnUL",    # ganti dengan access key kamu
    secret_key="SY9xregkSVUrf2O08lZpXsYzkhg1peB6r2yho4J7",    # ganti dengan secret key kamu
    secure=True                # True jika pakai https
)

BUCKET_NAME = "dokumen"

def upload_to_minio(image_data, content_type):
    filename = f"{uuid.uuid4()}.png"  # bisa pakai jpg/png tergantung mime
    image_io = io.BytesIO(image_data)
    image_size = len(image_data)

    minio_client.put_object(
        bucket_name=BUCKET_NAME,
        object_name=filename,
        data=image_io,
        length=image_size,
        content_type=content_type
    )

    return f"https://storage-api.sman16bekasi.id/{BUCKET_NAME}/{filename}"  # sesuaikan URL MinIO kamu

def extract_text_and_images(cell):
    html_parts = []
    nsmap = {
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    }

    def get_run_style(run):
        styles = []
        if run.bold:
            styles.append("font-weight:bold;")
        if run.italic:
            styles.append("font-style:italic;")
        if run.underline:
            styles.append("text-decoration:underline;")
        if run.font.size:
            try:
                styles.append(f"font-size:{run.font.size.pt}pt;")
            except:
                pass
        if run.font.name:
            styles.append(f"font-family:'{run.font.name}';")
        if run.font.color and run.font.color.rgb:
            styles.append(f"color:#{run.font.color.rgb};")
        return ''.join(styles)

    def extract_crop_from_run(run):
        src_rect = run._element.find('.//a:srcRect', namespaces=nsmap)
        if src_rect is not None:
            return {
                'l': int(src_rect.attrib.get('l', 0)),
                't': int(src_rect.attrib.get('t', 0)),
                'r': int(src_rect.attrib.get('r', 0)),
                'b': int(src_rect.attrib.get('b', 0)),
            }
        return None

    def crop_image_if_needed(image_data, crop):
        image = Image.open(io.BytesIO(image_data))
        width, height = image.size
        left = width * crop['l'] / 100000
        top = height * crop['t'] / 100000
        right = width * (1 - crop['r'] / 100000)
        bottom = height * (1 - crop['b'] / 100000)
        cropped = image.crop((left, top, right, bottom))
        buffer = io.BytesIO()
        cropped.save(buffer, format='PNG')
        buffer.seek(0)
        return buffer.read()

    def convert_omml_to_mathml(omml_element):
        try:
            xslt_root = etree.parse('OMML2MML.XSL')  # pastikan file ini tersedia
            transform = etree.XSLT(xslt_root)
            mathml = transform(omml_element)
            return etree.tostring(mathml, encoding='unicode')
        except Exception as e:
            print(f"OMML conversion error: {e}")
            return None

    # Parse equations first (outside of .paragraphs)
    math_elements = cell._element.findall('.//m:oMath', namespaces=nsmap)
    for math_elem in math_elements:
        mathml = convert_omml_to_mathml(math_elem)
        if mathml:
            html_parts.append(f"<p><math xmlns='http://www.w3.org/1998/Math/MathML'>{mathml}</math></p>")

    # Handle regular paragraphs and images
    for para in cell.paragraphs:
        para_html = ""
        for run in para.runs:
            run_text = run.text or ""
            style_attr = get_run_style(run)
            formatted = f"<span style='{style_attr}'>{run_text}</span>" if style_attr else run_text
            para_html += formatted

            # Handle images
            drawing_elements = run._element.findall('.//a:blip', namespaces=nsmap)
            for blip in drawing_elements:
                r_embed = blip.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if r_embed:
                    image_part = cell.part.related_parts[r_embed]
                    image_data = image_part.blob
                    crop = extract_crop_from_run(run)
                    if crop:
                        try:
                            image_data = crop_image_if_needed(image_data, crop)
                        except Exception as e:
                            print(f"Crop failed: {e} — using original image.")
                    image_url = upload_to_minio(image_data, image_part.content_type)
                    para_html += f"<br/><img src='{image_url}' style='max-width:100%; height:auto;' /><br/>"

        if para_html.strip():
            html_parts.append(f"<p>{para_html}</p>")

    return ''.join(html_parts)

@app.route("/parse-docx", methods=["POST"])
def parse_docx():
    file = request.files.get("file")
    if not file:
        return {"error": "No file"}, 400

    doc = Document(file)
    soal_list = []

    for i, row in enumerate(doc.tables[0].rows):
        if i == 0:
            continue
        cells = row.cells
        soal_data = {}
        for idx, key in enumerate(['soal', 'a', 'b', 'c', 'd', 'e', 'jawaban']):
            contents = extract_text_and_images(cells[idx])
            soal_data[key] = contents
        soal_list.append(soal_data)

    return jsonify(soal_list)

@app.route("/parse-docx/essay", methods=["POST"])
def parse_docx_essay():
    file = request.files.get("file")
    if not file:
        return {"error": "No file"}, 400

    doc = Document(file)
    soal_list = []

    for i, row in enumerate(doc.tables[0].rows):
        if i == 0:
            continue
        cells = row.cells
        soal_data = {}
        for idx, key in enumerate(['soal', 'jawaban']):
            contents = extract_text_and_images(cells[idx])
            soal_data[key] = contents
        soal_list.append(soal_data)

    return jsonify(soal_list)

if __name__ == "__main__":
    app.run(host='0.0.0.0', port=5000)
